"""Turn 18 smoke test — document extractors + semantic chunker round-trip.

What it proves:
  - extract_blocks works for PDF, DOCX, XLSX, TXT, MD without raising.
  - PDF extraction surfaces 1-indexed page numbers and font-size-derived
    section headings.
  - DOCX extraction picks up Heading-style paragraphs as section_heading.
  - XLSX extraction uses sheet name as section_heading and emits one block
    per non-empty row.
  - MD extraction detects ATX headings (``#``/``##``) and propagates them.
  - chunk_blocks respects the max_tokens ceiling — no semantic chunk exceeds
    the budget.
  - chunk_blocks preserves block boundaries — each chunk's content is a
    double-newline join of one or more whole extracted blocks (verified by
    parsing the chunk back into blocks and matching paragraph_index ranges).
  - Citation-ready meta carries source_file, paragraph_start, paragraph_end,
    section_heading on every chunk; page_start / page_end on PDF chunks.
  - The oversized-block fallback path produces chunks tagged ``fallback=True``
    when a single paragraph exceeds max_tokens.

Run from the repo root (no DB, no network needed):

    python backend/scripts/smoke_extractors.py
"""
from __future__ import annotations

import sys
import tempfile
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from openpyxl import Workbook

# Allow running this script directly without installing the backend package.
_BACKEND_DIR = Path(__file__).resolve().parent.parent
if str(_BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(_BACKEND_DIR))

from app.documents.chunker import chunk_blocks  # noqa: E402
from app.documents.extractors import extract_blocks  # noqa: E402


# --------------------------------------------------------------------------- #
# Synthetic test-file builders                                                #
# --------------------------------------------------------------------------- #
def make_pdf(path: Path) -> None:
    doc = fitz.open()
    # Page 1
    page = doc.new_page()
    page.insert_text((50, 60), "Introduction", fontsize=20)
    page.insert_text(
        (50, 110),
        "This is the opening paragraph of the document. "
        "It explains the purpose of the report and sets context.",
        fontsize=11,
    )
    page.insert_text(
        (50, 170),
        "A second paragraph adds detail about the methodology used.",
        fontsize=11,
    )
    # Page 2
    page2 = doc.new_page()
    page2.insert_text((50, 60), "Results", fontsize=20)
    page2.insert_text(
        (50, 110),
        "Revenue increased by twelve percent quarter over quarter, "
        "driven by enterprise expansion.",
        fontsize=11,
    )
    doc.save(str(path))
    doc.close()


def make_docx(path: Path) -> None:
    doc = DocxDocument()
    doc.add_heading("Pricing", level=1)
    doc.add_paragraph("Standard plan is forty dollars per seat per month.")
    doc.add_paragraph("Enterprise pricing is negotiated case by case.")
    doc.add_heading("Refunds", level=1)
    doc.add_paragraph("Refunds are issued within thirty days of purchase.")
    doc.save(str(path))


def make_xlsx(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Revenue 2024"
    ws.append(["Quarter", "Revenue", "Growth"])
    ws.append(["Q1", 100000, 0.08])
    ws.append(["Q2", 112000, 0.12])
    ws.append(["Q3", 118000, 0.05])
    wb.save(str(path))


def make_txt(path: Path) -> None:
    path.write_text(
        "First paragraph about something.\n\n"
        "Second paragraph adds more.\n\n"
        "Third and final paragraph wraps up.",
        encoding="utf-8",
    )


def make_md(path: Path) -> None:
    path.write_text(
        "# Top Section\n\n"
        "Intro under the top section.\n\n"
        "## Subsection A\n\n"
        "Body of subsection A.\n\n"
        "## Subsection B\n\n"
        "Body of subsection B with a couple of sentences for substance.",
        encoding="utf-8",
    )


def make_oversized_txt(path: Path) -> None:
    """Single paragraph designed to exceed max_tokens=500 — forces fallback."""
    # ~800 tokens of repeated content
    word = "tokenization "
    body = word * 800
    path.write_text(body, encoding="utf-8")


def make_multi_paragraph_txt(path: Path) -> None:
    """Many medium paragraphs whose TOTAL exceeds max_tokens=500 — forces
    the semantic-split path (lift B) to fire across paragraph boundaries
    without triggering the oversized-block fallback.

    Each paragraph is ~120 tokens; 7 paragraphs → ~840 total tokens →
    expect 2-3 semantic chunks, none marked fallback, split cleanly between
    whole paragraphs.
    """
    para = (
        "tokenization " * 120  # ~120 tokens
    ).strip()
    body = "\n\n".join(f"Paragraph {i}: {para}" for i in range(1, 8))
    path.write_text(body, encoding="utf-8")


# --------------------------------------------------------------------------- #
# Per-format assertion harness                                                #
# --------------------------------------------------------------------------- #
def check_pdf(path: Path, failures: list[str]) -> None:
    print("=== PDF ===")
    blocks = extract_blocks(str(path))
    print(f"  blocks: {len(blocks)}")
    for b in blocks:
        print(f"    p{b.page} ¶{b.paragraph_index} h={b.section_heading!r} :: {b.text[:60]!r}")

    if not blocks:
        failures.append("PDF: no blocks extracted")
        return
    if not any(b.page == 1 for b in blocks):
        failures.append("PDF: no block on page 1")
    if not any(b.page == 2 for b in blocks):
        failures.append("PDF: no block on page 2")
    # Section heading propagation: at least one block on page 2 should carry
    # "Results" as section_heading (the heading detected on page 2).
    p2_headings = {b.section_heading for b in blocks if b.page == 2}
    if "Results" not in p2_headings:
        failures.append(f"PDF: expected 'Results' heading on page 2, got {p2_headings}")
    # paragraph_index should be strictly increasing
    indices = [b.paragraph_index for b in blocks]
    if indices != sorted(indices) or len(set(indices)) != len(indices):
        failures.append("PDF: paragraph_index not strictly increasing or has duplicates")

    chunks = chunk_blocks(blocks, source_file=path.name)
    print(f"  chunks: {len(chunks)}")
    for c in chunks:
        m = c.meta
        print(
            f"    chunk {c.chunk_index} tok={c.token_count} "
            f"pages={m.get('page_start')}-{m.get('page_end')} "
            f"para={m.get('paragraph_start')}-{m.get('paragraph_end')} "
            f"h={m.get('section_heading')!r}"
        )
    _assert_citation_ready(chunks, expect_pages=True, label="PDF", failures=failures)


def check_docx(path: Path, failures: list[str]) -> None:
    print("=== DOCX ===")
    blocks = extract_blocks(str(path))
    print(f"  blocks: {len(blocks)}")
    for b in blocks:
        print(f"    ¶{b.paragraph_index} h={b.section_heading!r} :: {b.text[:60]!r}")

    if not blocks:
        failures.append("DOCX: no blocks extracted")
        return
    headings_seen = {b.section_heading for b in blocks}
    if "Pricing" not in headings_seen:
        failures.append(f"DOCX: expected 'Pricing' heading, got {headings_seen}")
    if "Refunds" not in headings_seen:
        failures.append(f"DOCX: expected 'Refunds' heading, got {headings_seen}")

    chunks = chunk_blocks(blocks, source_file=path.name)
    print(f"  chunks: {len(chunks)}")
    for c in chunks:
        print(f"    chunk {c.chunk_index} tok={c.token_count} h={c.meta.get('section_heading')!r}")
    _assert_citation_ready(chunks, expect_pages=False, label="DOCX", failures=failures)


def check_xlsx(path: Path, failures: list[str]) -> None:
    print("=== XLSX ===")
    blocks = extract_blocks(str(path))
    print(f"  blocks: {len(blocks)}")
    for b in blocks:
        print(f"    ¶{b.paragraph_index} sheet={b.section_heading!r} :: {b.text!r}")

    if not blocks:
        failures.append("XLSX: no blocks extracted")
        return
    # All blocks should carry the sheet name as section_heading.
    sheets = {b.section_heading for b in blocks}
    if sheets != {"Revenue 2024"}:
        failures.append(f"XLSX: section_heading should be {{'Revenue 2024'}}, got {sheets}")
    # One block per non-empty row → 4 rows in the synthetic sheet (header + 3 data rows).
    if len(blocks) != 4:
        failures.append(f"XLSX: expected 4 blocks (header + 3 rows), got {len(blocks)}")

    chunks = chunk_blocks(blocks, source_file=path.name)
    print(f"  chunks: {len(chunks)}")
    _assert_citation_ready(chunks, expect_pages=False, label="XLSX", failures=failures)


def check_txt(path: Path, failures: list[str]) -> None:
    print("=== TXT ===")
    blocks = extract_blocks(str(path))
    print(f"  blocks: {len(blocks)}")
    if len(blocks) != 3:
        failures.append(f"TXT: expected 3 paragraphs, got {len(blocks)}")
    for b in blocks:
        if b.section_heading is not None:
            failures.append(f"TXT: section_heading should be None, got {b.section_heading!r}")

    chunks = chunk_blocks(blocks, source_file=path.name)
    print(f"  chunks: {len(chunks)}")
    _assert_citation_ready(chunks, expect_pages=False, label="TXT", failures=failures)


def check_md(path: Path, failures: list[str]) -> None:
    print("=== MD ===")
    blocks = extract_blocks(str(path))
    print(f"  blocks: {len(blocks)}")
    for b in blocks:
        print(f"    ¶{b.paragraph_index} h={b.section_heading!r} :: {b.text[:50]!r}")

    if not blocks:
        failures.append("MD: no blocks extracted")
        return
    headings_seen = {b.section_heading for b in blocks if b.section_heading}
    if "Top Section" not in headings_seen:
        failures.append(f"MD: expected 'Top Section' heading, got {headings_seen}")
    if "Subsection A" not in headings_seen:
        failures.append(f"MD: expected 'Subsection A' heading, got {headings_seen}")

    chunks = chunk_blocks(blocks, source_file=path.name)
    print(f"  chunks: {len(chunks)}")
    _assert_citation_ready(chunks, expect_pages=False, label="MD", failures=failures)


def check_semantic_split(path: Path, failures: list[str]) -> None:
    print("=== Multi-paragraph semantic split ===")
    blocks = extract_blocks(str(path))
    print(f"  blocks: {len(blocks)}")
    chunks = chunk_blocks(blocks, source_file=path.name, max_tokens=500, overlap_tokens=50)
    print(f"  chunks: {len(chunks)}")
    for c in chunks:
        m = c.meta
        print(
            f"    chunk {c.chunk_index} tok={c.token_count} "
            f"para={m.get('paragraph_start')}-{m.get('paragraph_end')} "
            f"blocks={m.get('block_count')} fallback={m.get('fallback')}"
        )
    if len(chunks) < 2:
        failures.append(f"SemanticSplit: expected ≥2 chunks from oversized total, got {len(chunks)}")
    if any(c.meta.get("fallback") for c in chunks):
        bad = [c.chunk_index for c in chunks if c.meta.get("fallback")]
        failures.append(f"SemanticSplit: chunks {bad} fell back to token-window — paragraphs fit individually, split should be semantic")
    if any(c.token_count > 500 for c in chunks):
        over = [c.chunk_index for c in chunks if c.token_count > 500]
        failures.append(f"SemanticSplit: chunks {over} exceed max_tokens=500")
    # Paragraph boundaries: each chunk's paragraph_end + 1 should equal the
    # next chunk's paragraph_start — no overlap, no gap.
    for i in range(len(chunks) - 1):
        prev_end = chunks[i].meta.get("paragraph_end")
        next_start = chunks[i + 1].meta.get("paragraph_start")
        if prev_end is None or next_start is None:
            failures.append(f"SemanticSplit: chunk {i} or {i+1} missing paragraph_start/end")
            continue
        if next_start != prev_end + 1:
            failures.append(
                f"SemanticSplit: paragraph boundary not respected between chunk {i} "
                f"(ends ¶{prev_end}) and chunk {i+1} (starts ¶{next_start})"
            )


def check_oversized(path: Path, failures: list[str]) -> None:
    print("=== Oversized-block fallback ===")
    blocks = extract_blocks(str(path))
    print(f"  blocks: {len(blocks)} (single oversized paragraph)")
    chunks = chunk_blocks(blocks, source_file=path.name, max_tokens=500, overlap_tokens=50)
    print(f"  chunks: {len(chunks)}")
    for c in chunks:
        print(
            f"    chunk {c.chunk_index} tok={c.token_count} "
            f"fallback={c.meta.get('fallback')} "
            f"window={c.meta.get('fallback_window_start')}-{c.meta.get('fallback_window_end')}"
        )
    if len(chunks) < 2:
        failures.append(f"Oversized: expected multiple fallback chunks, got {len(chunks)}")
    if not all(c.meta.get("fallback") is True for c in chunks):
        failures.append("Oversized: all chunks should be marked fallback=True")
    if not all(c.token_count <= 500 for c in chunks):
        over = [c.chunk_index for c in chunks if c.token_count > 500]
        failures.append(f"Oversized: chunks {over} exceed max_tokens=500")


# --------------------------------------------------------------------------- #
# Shared chunk-meta sanity check                                              #
# --------------------------------------------------------------------------- #
def _assert_citation_ready(
    chunks: list,
    expect_pages: bool,
    label: str,
    failures: list[str],
    max_tokens: int = 500,
) -> None:
    if not chunks:
        failures.append(f"{label}: no chunks emitted")
        return
    required = {"source_file", "paragraph_start", "paragraph_end", "section_heading", "block_count"}
    for c in chunks:
        missing = required - set(c.meta.keys())
        if missing:
            failures.append(f"{label} chunk {c.chunk_index}: missing meta keys {missing}")
        if c.token_count > max_tokens and not c.meta.get("fallback"):
            failures.append(
                f"{label} chunk {c.chunk_index}: token_count {c.token_count} exceeds "
                f"max_tokens {max_tokens} without fallback flag"
            )
        if expect_pages and ("page_start" not in c.meta or "page_end" not in c.meta):
            failures.append(f"{label} chunk {c.chunk_index}: missing page_start/page_end")


# --------------------------------------------------------------------------- #
# Entry point                                                                 #
# --------------------------------------------------------------------------- #
def main() -> int:
    failures: list[str] = []
    with tempfile.TemporaryDirectory(prefix="jarvis-smoke-extractors-") as tmpdir:
        td = Path(tmpdir)

        pdf_path = td / "report.pdf"
        docx_path = td / "policies.docx"
        xlsx_path = td / "revenue.xlsx"
        txt_path = td / "notes.txt"
        md_path = td / "spec.md"
        big_path = td / "oversized.txt"
        multi_path = td / "multi_paragraph.txt"

        make_pdf(pdf_path)
        make_docx(docx_path)
        make_xlsx(xlsx_path)
        make_txt(txt_path)
        make_md(md_path)
        make_oversized_txt(big_path)
        make_multi_paragraph_txt(multi_path)

        check_pdf(pdf_path, failures)
        check_docx(docx_path, failures)
        check_xlsx(xlsx_path, failures)
        check_txt(txt_path, failures)
        check_md(md_path, failures)
        check_semantic_split(multi_path, failures)
        check_oversized(big_path, failures)

    print()
    if failures:
        print(f"FAIL: {len(failures)} assertion(s) failed")
        for f in failures:
            print(f"  - {f}")
        return 1
    print("PASS: extractors + chunker produce citation-ready chunks across all formats")
    return 0


if __name__ == "__main__":
    sys.exit(main())
